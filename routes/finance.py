"""维修做账 - 发票录入→清单生成→文档输出"""
import json
import random
from datetime import datetime, date
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, g, send_file
from flask_login import login_required, current_user
from models import db, FinanceBatch, FinanceInvoice, FinanceDraft, FinanceDraftPart, PartPrice, WorkOrder, Asset, log_audit, can_access, AcceptanceTemplate
from sqlalchemy import func
from fuzzywuzzy import fuzz
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

finance_bp = Blueprint('finance', __name__, url_prefix='/finance')


def check_access():
    if not can_access('维修做账'):
        flash('无权访问维修做账', 'danger')
        return False
    return True


# WorkOrder device_type → Asset device_type 映射
WO_TO_ASSET_TYPE = {
    '电脑': ['PC'],
    '硬件': ['PC'],
    '软件': ['PC'],
    '键盘': ['PC'],
    '鼠标': ['PC'],
    '打印机': ['激光打印机', '针式打印机'],
}


def find_best_asset(hospital_id, dept, device_type_str, building='', floor=''):
    """模糊匹配最合适的资产，按相似度取最优，找不到返回 None"""
    target_types = WO_TO_ASSET_TYPE.get(device_type_str, None)

    if not dept or not target_types:
        return None

    best_asset = None
    best_score = 0

    # 收集所有候选资产（同设备类型），模糊匹配科室名
    for dt in target_types:
        candidates = Asset.query.filter(
            Asset.hospital_id == hospital_id,
            Asset.device_type == dt,
        ).all()
        for a in candidates:
            if not a.department:
                continue
            score = max(fuzz.token_sort_ratio(dept, a.department), fuzz.partial_ratio(dept, a.department))
            if score > best_score:
                best_score = score
                best_asset = a

    # 阈值 50 以上才接受
    if best_asset and best_score >= 70:
        return best_asset

    # 仍找不到 → 退到楼栋+楼层+设备类型
    if building:
        for dt in target_types:
            asset = Asset.query.filter(
                Asset.hospital_id == hospital_id,
                Asset.building == building,
                Asset.floor == floor,
                Asset.device_type == dt,
            ).first()
            if asset:
                return asset

    # 同楼栋（放宽楼层）
    if building:
        for dt in target_types:
            asset = Asset.query.filter(
                Asset.hospital_id == hospital_id,
                Asset.building == building,
                Asset.device_type == dt,
            ).first()
            if asset:
                return asset

    # 还找不到 → 用"通用"科室的同类型资产
    for dt in target_types:
        asset = Asset.query.filter(
            Asset.hospital_id == hospital_id,
            Asset.department == '通用',
            Asset.device_type == dt,
        ).first()
        if asset:
            return asset

    return None


# ==================== 列表 ====================

@finance_bp.route('/')
@login_required
def index():
    if not check_access():
        return redirect(url_for('main.dashboard'))
    batches = FinanceBatch.query.order_by(FinanceBatch.created_at.desc()).all()
    return render_template('finance/list.html', batches=batches)


# ==================== 新建批次 ====================

@finance_bp.route('/create', methods=['GET', 'POST'])
@login_required
def create():
    if not check_access():
        return redirect(url_for('main.dashboard'))
    if request.method == 'POST':
        company = request.form.get('company_name', '').strip()
        payee = request.form.get('payee', '').strip()
        bank = request.form.get('bank_name', '').strip()
        account = request.form.get('bank_account', '').strip()

        # 解析发票 JSON
        invoices_data = request.form.get('invoices_json', '[]')
        try:
            invoices = json.loads(invoices_data)
        except json.JSONDecodeError:
            invoices = []

        total = sum(float(inv.get('amount', 0)) for inv in invoices)

        batch = FinanceBatch(
            company_name=company,
            payee=payee,
            bank_name=bank,
            bank_account=account,
            total_amount=total,
            status='draft',
            created_by=current_user.display_name or current_user.username,
            hospital_id=getattr(g, 'hospital_id', 1) or 1,
        )
        db.session.add(batch)
        db.session.flush()

        for inv in invoices:
            fi = FinanceInvoice(
                batch_id=batch.id,
                invoice_no=inv.get('invoice_no', ''),
                amount=float(inv.get('amount', 0)),
            )
            db.session.add(fi)

        db.session.commit()
        log_audit('create', 'finance_batch', current_user.display_name or current_user.username,
                  f'创建维修做账批次 #{batch.id}: {company}')
        flash('发票批次已创建', 'success')
        return redirect(url_for('finance.index'))

    return render_template('finance/create.html')


# ==================== 详情 ====================

@finance_bp.route('/<int:batch_id>')
@login_required
def detail(batch_id):
    if not check_access():
        return redirect(url_for('main.dashboard'))
    batch = FinanceBatch.query.get_or_404(batch_id)
    invoices = batch.invoices.all()
    drafts = batch.drafts.order_by(FinanceDraft.sort_order).all()
    return render_template('finance/detail.html', batch=batch, invoices=invoices, drafts=drafts)


# ==================== 生成维修清单草稿 ====================

@finance_bp.route('/<int:batch_id>/generate_drafts', methods=['POST'])
@login_required
def generate_drafts(batch_id):
    """根据选中月份的已完成工单自动生成维修清单草稿"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403

    batch = FinanceBatch.query.get_or_404(batch_id)
    target_amount = float(batch.total_amount)
    hid = getattr(g, 'hospital_id', 1) or 1

    # 获取选中的月份（默认当月）
    year = request.form.get('year', datetime.now().year, type=int)
    month = request.form.get('month', datetime.now().month, type=int)
    month_start = datetime(year, month, 1)
    if month == 12:
        month_end = datetime(year + 1, 1, 1)
    else:
        month_end = datetime(year, month + 1, 1)

    orders = WorkOrder.query.filter(
        WorkOrder.hospital_id == hid,
        WorkOrder.status == 'completed',
        WorkOrder.created_at >= month_start,
        WorkOrder.created_at < month_end,
    ).all()

    if not orders:
        flash(f'{year}年{month}月暂无已完成硬件报修工单，无法生成清单', 'warning')
        return redirect(url_for('finance.detail', batch_id=batch_id))

    # 获取零件列表，按类型分类
    all_parts = PartPrice.query.filter(
        PartPrice.hospital_id == hid,
        PartPrice.unit_price > 0,
    ).all()
    if not all_parts:
        flash('零件价格库为空，请先添加零件', 'danger')
        return redirect(url_for('finance.detail', batch_id=batch_id))
    from collections import defaultdict
    parts_by_type = defaultdict(list)
    for p in all_parts:
        parts_by_type[p.category].append(p)

    # 删除旧草稿
    FinanceDraftPart.query.filter(
        FinanceDraftPart.draft_id.in_(
            db.session.query(FinanceDraft.id).filter(FinanceDraft.batch_id == batch_id)
        )
    ).delete(synchronize_session=False)
    FinanceDraft.query.filter_by(batch_id=batch_id).delete()
    db.session.flush()

    # 逐条工单生成维修记录
    sort = 0
    remaining = target_amount
    for o in orders:
        if remaining <= 0:
            break

        device = '打印机' if o.device_type in ('PR', '打印机', '激光打印机', '针式打印机') else '电脑'
        dept = o.department or '未知科室'

        # 模糊匹配资产
        asset = find_best_asset(hid, dept, o.device_type, o.building, o.floor)

        report_date = o.created_at

        # 按设备类型匹配零件池
        if device == '电脑':
            pool = parts_by_type.get('电脑', [])
        else:
            pool = parts_by_type.get('激光打印机', []) + parts_by_type.get('针式打印机', [])
        pool = [p for p in pool if float(p.unit_price) <= remaining]

        # 随机选配件：最多3种
        random.shuffle(pool)
        chosen_parts = []
        draft_total = 0
        max_parts = min(3, len(pool))
        for p in pool[:max_parts * 3]:
            if len(chosen_parts) >= max_parts:
                break
            if p.product_name in [cp.product_name for cp in chosen_parts]:
                continue
            price = float(p.unit_price)
            chosen_parts.append(p)
            draft_total += price

        if not chosen_parts:
            continue

        content = ' '.join(p.product_name for p in chosen_parts)

        draft = FinanceDraft(
            batch_id=batch_id,
            asset_id=asset.id if asset else None,
            device_type=device,
            department=dept,
            report_date=report_date,
            repair_content=content,
            total_amount=draft_total,
            sort_order=sort,
            status='draft',
        )
        db.session.add(draft)
        db.session.flush()

        for p in chosen_parts:
            dp = FinanceDraftPart(
                draft_id=draft.id,
                part_name=p.product_name,
                unit=p.unit,
                quantity=1,
                unit_price=p.unit_price,
                amount=p.unit_price,
            )
            db.session.add(dp)

        remaining -= draft_total
        sort += 1

    batch.status = 'generated'
    db.session.commit()

    flash(f'已生成 {sort} 条维修清单草稿，合计 ¥{target_amount - remaining:.2f} / ¥{target_amount:.2f}', 'success')
    return redirect(url_for('finance.detail', batch_id=batch_id))


# ==================== 编辑草稿（AJAX） ====================

@finance_bp.route('/draft/<int:draft_id>/save', methods=['POST'])
@login_required
def save_draft(draft_id):
    draft = FinanceDraft.query.get_or_404(draft_id)
    draft.repair_content = request.form.get('repair_content', draft.repair_content)
    draft.total_amount = request.form.get('total_amount', 0, type=float)
    db.session.commit()
    return jsonify({'ok': True})


@finance_bp.route('/draft/<int:draft_id>/part/add', methods=['POST'])
@login_required
def add_part(draft_id):
    draft = FinanceDraft.query.get_or_404(draft_id)
    name = request.form.get('part_name', '').strip()
    price = request.form.get('unit_price', 0, type=float)
    dp = FinanceDraftPart(
        draft_id=draft_id,
        part_name=name,
        unit=request.form.get('unit', '个'),
        quantity=1,
        unit_price=price,
        amount=price,
    )
    db.session.add(dp)
    draft.total_amount = float(draft.total_amount or 0) + price
    db.session.commit()
    return jsonify({'ok': True, 'id': dp.id, 'total': float(draft.total_amount)})


@finance_bp.route('/draft/<int:draft_id>/part/<int:part_id>/delete', methods=['POST'])
@login_required
def delete_draft_part(draft_id, part_id):
    dp = FinanceDraftPart.query.get_or_404(part_id)
    draft = FinanceDraft.query.get_or_404(draft_id)
    draft.total_amount = max(0, float(draft.total_amount or 0) - float(dp.amount or 0))
    db.session.delete(dp)
    db.session.commit()
    return jsonify({'ok': True, 'total': float(draft.total_amount)})


@finance_bp.route('/draft/<int:draft_id>/delete', methods=['POST'])
@login_required
def delete_draft(draft_id):
    draft = FinanceDraft.query.get_or_404(draft_id)
    batch_id = draft.batch_id
    FinanceDraftPart.query.filter_by(draft_id=draft_id).delete()
    db.session.delete(draft)
    db.session.commit()
    flash('维修记录已删除', 'success')
    return redirect(url_for('finance.detail', batch_id=batch_id))


# ==================== 选择资产 (AJAX) ====================

@finance_bp.route('/draft/<int:draft_id>/select_asset', methods=['POST'])
@login_required
def select_asset(draft_id):
    """手动选择/更换草稿关联的资产"""
    draft = FinanceDraft.query.get_or_404(draft_id)
    asset_id = request.form.get('asset_id', type=int)
    if asset_id:
        asset = db.session.get(Asset, asset_id)
        if asset:
            draft.asset_id = asset.id
            db.session.commit()
            return jsonify({'ok': True, 'asset_no': asset.asset_no, 'brand': asset.brand, 'model': asset.model_no})
    else:
        draft.asset_id = None
        db.session.commit()
        return jsonify({'ok': True, 'asset_no': None})
    return jsonify({'ok': False, 'error': '资产不存在'}), 404


@finance_bp.route('/assets_by_dept')
@login_required
def assets_by_dept():
    """按科室模糊搜索资产"""
    q = request.args.get('q', '').strip()
    hid = getattr(g, 'hospital_id', 1) or 1
    query = Asset.query.filter(Asset.hospital_id == hid)
    if q:
        query = query.filter(
            db.or_(
                Asset.department.ilike(f'%{q}%'),
                Asset.asset_no.ilike(f'%{q}%'),
                Asset.brand.ilike(f'%{q}%'),
                Asset.model_no.ilike(f'%{q}%'),
            )
        )
    assets = query.order_by(Asset.department, Asset.device_type).limit(200).all()
    data = [{
        'id': a.id,
        'asset_no': a.asset_no,
        'device_type': a.device_type,
        'brand': a.brand or '',
        'model_no': a.model_no or '',
        'department': a.department or '',
    } for a in assets]
    return jsonify(data)


# ==================== 生成文档 ====================

@finance_bp.route('/<int:batch_id>/generate_docs', methods=['POST'])
@login_required
def generate_docs(batch_id):
    """标记所有清单为已生成"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    drafts = FinanceDraft.query.filter_by(batch_id=batch_id).all()
    for d in drafts:
        d.status = 'generated'
    batch = FinanceBatch.query.get_or_404(batch_id)
    batch.status = 'generated'
    db.session.commit()
    flash('文档已生成', 'success')
    return redirect(url_for('finance.detail', batch_id=batch_id))


@finance_bp.route('/<int:batch_id>/print_acceptance')
@login_required
def print_acceptance(batch_id):
    """打印验收单（2×2 A4排版）"""
    if not check_access():
        return redirect(url_for('main.dashboard'))
    batch = FinanceBatch.query.get_or_404(batch_id)
    drafts = batch.drafts.order_by(FinanceDraft.sort_order).all()
    return render_template('finance/print_acceptance.html', batch=batch, drafts=drafts)


@finance_bp.route('/<int:batch_id>/export_acceptance_excel')
@login_required
def export_acceptance_excel(batch_id):
    """导出验收单为 Excel（A4 2×2 排版）"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    import io
    batch = FinanceBatch.query.get_or_404(batch_id)
    drafts = batch.drafts.order_by(FinanceDraft.sort_order).all()
    wb = openpyxl.Workbook()
    total = len(drafts)
    per_page = 8
    pages = (total + per_page - 1) // per_page

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    for page in range(pages):
        if page == 0:
            ws = wb.active
            ws.title = f'验收单-1'
        else:
            ws = wb.create_sheet(title=f'验收单-{page+1}')

        ws.page_setup.orientation = 'portrait'
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.page_setup.fitToPage = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        ws.sheet_properties.pageSetUpPr = openpyxl.worksheet.properties.PageSetupProperties(fitToPage=True)

        # 列：左半页(5列) + 间隙 + 右半页(5列)
        col_widths = [7, 14, 4, 7, 14, 2, 7, 14, 4, 7, 14]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w * 1.6

        # 行高紧凑
        for r in range(1, 40):
            ws.row_dimensions[r].height = 11

        for card_idx in range(per_page):
            idx = page * per_page + card_idx
            if idx >= total:
                break
            d = drafts[idx]

            # 2列 × 4行：每张卡片占7行
            row = (card_idx // 2) * 8 + 1  # 每张卡片7行 + 1行间隙
            col = 1 if card_idx % 2 == 0 else 7

            # 标题
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+4)
            c = ws.cell(row=row, column=col, value='维修验收单')
            c.font = Font(size=10, bold=True)
            c.alignment = Alignment(horizontal='center', vertical='center')

            # 日期 + 科室
            r = row + 1
            ws.cell(row=r, column=col, value='日期：').font = Font(bold=True, size=7)
            ws.merge_cells(start_row=r, start_column=col+1, end_row=r, end_column=col+1)
            ws.cell(row=r, column=col+1, value=d.report_date.strftime('%m-%d') if d.report_date else '').font = Font(size=7)
            ws.cell(row=r, column=col+2, value='科室：').font = Font(bold=True, size=7)
            ws.merge_cells(start_row=r, start_column=col+3, end_row=r, end_column=col+4)
            ws.cell(row=r, column=col+3, value=d.department or '').font = Font(size=7)

            # 设备 + 编号
            r = row + 2
            ws.cell(row=r, column=col, value='设备：').font = Font(bold=True, size=7)
            ws.merge_cells(start_row=r, start_column=col+1, end_row=r, end_column=col+1)
            ws.cell(row=r, column=col+1, value=d.device_type or '').font = Font(size=7)
            ws.cell(row=r, column=col+2, value='编号：').font = Font(bold=True, size=7)
            ws.merge_cells(start_row=r, start_column=col+3, end_row=r, end_column=col+4)
            ws.cell(row=r, column=col+3, value=d.asset.asset_no if d.asset else '').font = Font(size=7)

            # 内容
            r = row + 3
            ws.cell(row=r, column=col, value='内容：').font = Font(bold=True, size=7)
            ws.merge_cells(start_row=r, start_column=col+1, end_row=r+1, end_column=col+4)
            cc = ws.cell(row=r, column=col+1, value=d.repair_content or '')
            cc.font = Font(size=7)
            cc.alignment = Alignment(wrap_text=True, vertical='top')

            # 金额 + 完成状态
            r = row + 5
            ws.cell(row=r, column=col, value='金额：').font = Font(bold=True, size=7)
            ws.cell(row=r, column=col+1, value=f'¥{d.total_amount:.0f}').font = Font(size=8, bold=True)
            ws.merge_cells(start_row=r, start_column=col+2, end_row=r, end_column=col+4)
            ws.cell(row=r, column=col+2, value='☑已完成 □未完成').font = Font(size=7)
            ws.cell(row=r, column=col+2).alignment = Alignment(horizontal='center')

            # 签字
            r = row + 6
            ws.cell(row=r, column=col, value='验收人：').font = Font(bold=True, size=7)
            ws.cell(row=r, column=col+3, value='日期：').font = Font(bold=True, size=7)

            # 外框
            for br in range(row, row+7):
                for bc in range(col, col+5):
                    c = ws.cell(row=br, column=bc)
                    c.border = Border(
                        left=Side(style='thin') if bc == col else Side(style='hair'),
                        right=Side(style='thin') if bc == col+4 else Side(style='hair'),
                        top=Side(style='thin') if br == row else Side(style='hair'),
                        bottom=Side(style='thin') if br == row+6 else Side(style='hair')
                    )

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'验收单_{batch.id}_{datetime.now().strftime("%Y%m%d")}.xlsx'
    )


# ==================== 删除批次 ====================

@finance_bp.route('/<int:batch_id>/delete', methods=['POST'])
@login_required
def delete_batch(batch_id):
    batch = FinanceBatch.query.get_or_404(batch_id)
    # 级联删除
    FinanceDraftPart.query.filter(
        FinanceDraftPart.draft_id.in_(
            db.session.query(FinanceDraft.id).filter(FinanceDraft.batch_id == batch_id)
        )
    ).delete(synchronize_session=False)
    FinanceDraft.query.filter_by(batch_id=batch_id).delete()
    FinanceInvoice.query.filter_by(batch_id=batch_id).delete()
    db.session.delete(batch)
    db.session.commit()
    flash('批次已删除', 'success')
    return redirect(url_for('finance.index'))

# ==================== 导出Excel ====================

@finance_bp.route('/<int:batch_id>/export')
@login_required
def export_excel(batch_id):
    """导出维修做账的4份文档为Excel"""
    from flask import send_file
    batch = FinanceBatch.query.get_or_404(batch_id)
    drafts = batch.drafts.order_by(FinanceDraft.sort_order).all()
    invoices = batch.invoices.all()

    wb = Workbook()

    # 样式
    header_font = Font(bold=True, size=11)
    title_font = Font(bold=True, size=14)
    center = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )
    header_fill = PatternFill(start_color='F1F5F9', end_color='F1F5F9', fill_type='solid')

    # ====== Sheet 1: 维修清单 ======
    ws1 = wb.active
    ws1.title = '维修清单'
    headers1 = ['日期', '资产编码', '类型', '维修内容', '合计金额', '科室']
    ws1.append(headers1)
    for col in range(1, 7):
        cell = ws1.cell(row=1, column=col)
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border
        cell.fill = header_fill

    total = 0
    for d in drafts:
        ws1.append([
            d.report_date.strftime('%Y-%m-%d') if d.report_date else '',
            d.asset.asset_no if d.asset else '',
            d.device_type,
            d.repair_content,
            float(d.total_amount),
            d.department,
        ])
        total += float(d.total_amount)
        for col in range(1, 7):
            ws1.cell(row=ws1.max_row, column=col).border = thin_border

    ws1.append(['', '', '', '合计：', total, ''])
    ws1.cell(row=ws1.max_row, column=4).font = Font(bold=True)
    ws1.cell(row=ws1.max_row, column=5).font = Font(bold=True)

    ws1.column_dimensions['A'].width = 14
    ws1.column_dimensions['B'].width = 20
    ws1.column_dimensions['C'].width = 10
    ws1.column_dimensions['D'].width = 40
    ws1.column_dimensions['E'].width = 12
    ws1.column_dimensions['F'].width = 14

    # ====== Sheet 2: 维修小结（汇总一行） ======
    ws2 = wb.create_sheet('维修小结')
    headers2 = ['日期', '维修费', '数量', '金额']
    ws2.append(headers2)
    for col in range(1, 5):
        cell = ws2.cell(row=1, column=col)
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border
        cell.fill = header_fill

    total_count = len(drafts)
    total_amount = sum(float(d.total_amount) for d in drafts)
    batch_date = batch.created_at.strftime('%Y-%m-%d') if batch.created_at else datetime.now().strftime('%Y-%m-%d')
    ws2.append([batch_date, '维修费', total_count, total_amount])
    for col in range(1, 5):
        ws2.cell(row=2, column=col).border = thin_border
    ws2.cell(row=2, column=4).number_format = '#,##0.00'

    ws2.column_dimensions['A'].width = 14
    ws2.column_dimensions['B'].width = 40
    ws2.column_dimensions['C'].width = 8
    ws2.column_dimensions['D'].width = 14

    # ====== Sheet 3: 验收单（A4排版，每页4份，2×2排列，带边框） ======
    ws3 = wb.create_sheet('验收单')

    # A4 页面设置
    ws3.page_setup.paperSize = ws3.PAPERSIZE_A4
    ws3.page_setup.orientation = 'portrait'
    ws3.page_setup.fitToWidth = 1
    ws3.page_setup.fitToHeight = 0

    # 列宽：左份(A-D) + 间隔(E) + 右份(F-I)
    for col, w in {1: 12, 2: 20, 3: 12, 4: 20, 5: 3, 6: 12, 7: 20, 8: 12, 9: 20}.items():
        ws3.column_dimensions[chr(64 + col)].width = w

    def _write_form(ws, d, r, c):
        """在(r,c)位置写入一份验收单，加边框"""
        # 标题
        cell = ws.cell(row=r, column=c, value='维修验收单')
        cell.font = title_font
        cell.alignment = center
        ws.merge_cells(start_row=r, start_column=c, end_row=r, end_column=c+3)

        # 日期 + 科室
        ws.cell(row=r+1, column=c, value='维修日期').font = Font(bold=True)
        ws.cell(row=r+1, column=c+1, value=d.report_date.strftime('%Y-%m-%d') if d.report_date else '')
        ws.cell(row=r+1, column=c+2, value='维修科室').font = Font(bold=True)
        ws.cell(row=r+1, column=c+3, value=d.department)

        # 维修物品（跨3列）：有资产则显品牌型号，没有则显示设备类型
        if d.asset:
            brand_model = f"{d.asset.brand or ''} {d.asset.model_no or d.asset.asset_no or ''}".strip()
        else:
            brand_model = d.device_type or '维修'
        ws.cell(row=r+2, column=c, value='维修物品').font = Font(bold=True)
        ws.cell(row=r+2, column=c+1, value=brand_model)
        ws.merge_cells(start_row=r+2, start_column=c+1, end_row=r+2, end_column=c+3)

        # 维修内容（跨3列）
        ws.cell(row=r+3, column=c, value='维修内容').font = Font(bold=True)
        ws.cell(row=r+3, column=c+1, value=d.repair_content)
        ws.merge_cells(start_row=r+3, start_column=c+1, end_row=r+3, end_column=c+3)

        # 金额
        ws.cell(row=r+4, column=c, value='金额').font = Font(bold=True)
        ws.cell(row=r+4, column=c+1, value=f'¥{d.total_amount:.2f}')

        # 是否已完成
        ws.cell(row=r+5, column=c, value='是否已完成').font = Font(bold=True)
        ws.cell(row=r+5, column=c+1, value='☑是       □否')

        # 验收人签字
        ws.cell(row=r+6, column=c, value='验收人签字').font = Font(bold=True)

        # 全表单加细边框
        for dr in range(7):
            for dc in range(4):
                ws.cell(row=r+dr, column=c+dc).border = thin_border

    for i, d in enumerate(drafts):
        pos = i % 4          # 0=左上，1=右上，2=左下，3=右下
        page = i // 4        # 第几页
        base_row = page * 16 # 每页16行（7+1+7+1）
        col = 1 if pos % 2 == 0 else 6
        row = base_row + 1 if pos < 2 else base_row + 9
        _write_form(ws3, d, row, col)

    # 行高统一设为22磅
    for r in range(1, max(len(drafts) * 4 + 1, 10)):
        ws3.row_dimensions[r].height = 22

    # ====== Sheet 4: 发票汇总 ======
    ws4 = wb.create_sheet('发票汇总')
    headers4 = ['序号', '收款单位', '开户银行', '账号', '发票号码', '发票总金额']
    ws4.append(headers4)
    for col in range(1, 7):
        cell = ws4.cell(row=1, column=col)
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border
        cell.fill = header_fill

    for idx, inv in enumerate(invoices):
        ws4.append([
            idx + 1,
            batch.payee or '',
            batch.bank_name or '',
            batch.bank_account or '',
            inv.invoice_no,
            float(inv.amount),
        ])
        for col in range(1, 7):
            ws4.cell(row=ws4.max_row, column=col).border = thin_border

    ws4.append(['', '', '', '', '总计：', float(batch.total_amount)])
    ws4.cell(row=ws4.max_row, column=5).font = Font(bold=True)
    ws4.cell(row=ws4.max_row, column=6).font = Font(bold=True)

    ws4.column_dimensions['A'].width = 8
    ws4.column_dimensions['B'].width = 20
    ws4.column_dimensions['C'].width = 20
    ws4.column_dimensions['D'].width = 25
    ws4.column_dimensions['E'].width = 18
    ws4.column_dimensions['F'].width = 14

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    company = batch.company_name.replace('/', '_').replace('\\', '_') if batch.company_name else '维修做账'
    fname = f'{company}_维修做账_{datetime.now().strftime("%Y%m%d")}.xlsx'
    return send_file(output, as_attachment=True, download_name=fname, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


# ==================== 验收单模板在线设计器 ====================

ACCEPTANCE_FIELD_DEFS = [
    {'key': 'date', 'label': '维修日期', 'default': '2024-01-15', 'type': 'date'},
    {'key': 'department', 'label': '科室', 'default': '信息科', 'type': 'text'},
    {'key': 'device_type', 'label': '设备类型', 'default': '打印机', 'type': 'text'},
    {'key': 'device_brand', 'label': '设备品牌', 'default': 'HP', 'type': 'text'},
    {'key': 'device_model', 'label': '设备型号', 'default': 'LaserJet Pro', 'type': 'text'},
    {'key': 'asset_no', 'label': '资产编号', 'default': 'ZC-2024-001', 'type': 'text'},
    {'key': 'repair_content', 'label': '维修内容', 'default': '更换硒鼓、清洁进纸轮', 'type': 'text'},
    {'key': 'total_amount', 'label': '金额(元)', 'default': '580.00', 'type': 'number'},
    {'key': 'complete_status', 'label': '完成情况', 'default': '☑已完成  □未完成', 'type': 'text'},
    {'key': 'acceptor', 'label': '验收人签字', 'default': '', 'type': 'sign'},
    {'key': 'accept_date', 'label': '验收日期', 'default': '2024-01-20', 'type': 'date'},
    {'key': 'remark', 'label': '备注', 'default': '', 'type': 'text'},
    {'key': 'work_order_no', 'label': '工单编号', 'default': 'GD-2024-001', 'type': 'text'},
    {'key': 'repair_person', 'label': '维修人员', 'default': '张三', 'type': 'text'},
    {'key': 'section_title', 'label': '——— 分隔线 ———', 'default': '', 'type': 'separator'},
]


@finance_bp.route('/acceptance_designer')
@login_required
def acceptance_designer():
    """验收单在线设计器首页"""
    if not check_access():
        return redirect(url_for('main.dashboard'))
    templates = AcceptanceTemplate.query.order_by(AcceptanceTemplate.is_default.desc(), AcceptanceTemplate.id.desc()).all()
    return render_template('finance/acceptance_designer.html', templates=templates,
                           field_defs=ACCEPTANCE_FIELD_DEFS)


@finance_bp.route('/acceptance_designer/<int:tpl_id>')
@login_required
def acceptance_designer_edit(tpl_id):
    """编辑验收单模板"""
    if not check_access():
        return redirect(url_for('main.dashboard'))
    tpl = AcceptanceTemplate.query.get_or_404(tpl_id)
    templates = AcceptanceTemplate.query.order_by(AcceptanceTemplate.is_default.desc(), AcceptanceTemplate.id.desc()).all()
    return render_template('finance/acceptance_designer.html', templates=templates,
                           field_defs=ACCEPTANCE_FIELD_DEFS, current_tpl=tpl)


@finance_bp.route('/acceptance_designer/create', methods=['POST'])
@login_required
def acceptance_designer_create():
    """新建模板"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    name = request.form.get('name', '新模板').strip()
    import json
    layout = {
        'cols': 12, 'rows': 24,
        'page_width': 196, 'page_height': 277,
        'fields': [],
        'title': '维修验收单',
        'title_font_size': 12,
        'border_style': 'solid',
        'border_width': 1
    }
    tpl = AcceptanceTemplate(name=name, layout_json=json.dumps(layout, ensure_ascii=False))
    db.session.add(tpl)
    db.session.commit()
    return jsonify({'ok': True, 'id': tpl.id})


@finance_bp.route('/acceptance_designer/<int:tpl_id>/save', methods=['POST'])
@login_required
def acceptance_designer_save(tpl_id):
    """保存模板布局"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    tpl = AcceptanceTemplate.query.get_or_404(tpl_id)
    data = request.get_json(force=True) or {}
    if 'layout' in data:
        tpl.set_layout(data['layout'])
    if 'name' in data:
        tpl.name = data['name']
    if 'is_default' in data:
        # 唯一默认
        if data['is_default']:
            AcceptanceTemplate.query.filter(AcceptanceTemplate.id != tpl_id).update({'is_default': False})
        tpl.is_default = data['is_default']
    db.session.commit()
    return jsonify({'ok': True, 'id': tpl.id})


@finance_bp.route('/acceptance_designer/<int:tpl_id>/delete', methods=['POST'])
@login_required
def acceptance_designer_delete(tpl_id):
    """删除模板"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    tpl = AcceptanceTemplate.query.get_or_404(tpl_id)
    db.session.delete(tpl)
    db.session.commit()
    return jsonify({'ok': True})


@finance_bp.route('/acceptance_designer/<int:tpl_id>/preview')
@login_required
def acceptance_designer_preview(tpl_id):
    """预览验收单模板效果（用模拟数据）"""
    if not check_access():
        return redirect(url_for('main.dashboard'))
    tpl = AcceptanceTemplate.query.get_or_404(tpl_id)
    layout = tpl.get_layout()
    # 模拟数据
    mock_data = {
        'date': '2024-01-15',
        'department': '信息科',
        'device_type': '打印机',
        'device_brand': 'HP',
        'device_model': 'LaserJet Pro',
        'asset_no': 'ZC-2024-001',
        'repair_content': '更换硒鼓、清洁进纸轮、调整打印位置',
        'total_amount': '580.00',
        'complete_status': '☑已完成  □未完成',
        'acceptor': '',
        'accept_date': '2024-01-20',
        'remark': '',
        'work_order_no': 'GD-2024-001',
        'repair_person': '张三',
    }
    return render_template('finance/acceptance_render.html', layout=layout, data=mock_data, is_preview=True)


@finance_bp.route('/<int:batch_id>/acceptance_preview')
@finance_bp.route('/<int:batch_id>/acceptance_render')
@login_required
def batch_acceptance_render(batch_id):
    """用模板渲染批次的验收单"""
    if not check_access():
        return redirect(url_for('main.dashboard'))
    batch = FinanceBatch.query.get_or_404(batch_id)
    # 获取默认模板
    tpl = AcceptanceTemplate.query.filter_by(is_default=True).first()
    if not tpl:
        tpl = AcceptanceTemplate.query.first()
    if not tpl:
        flash('请先创建验收单模板', 'warning')
        return redirect(url_for('finance.acceptance_designer'))
    layout = tpl.get_layout()
    drafts = batch.drafts.order_by(FinanceDraft.sort_order).all()
    return render_template('finance/acceptance_render.html', layout=layout, drafts=drafts,
                           batch=batch, is_preview=False)


@finance_bp.route('/acceptance_designer/import_excel', methods=['POST'])
@login_required
def acceptance_designer_import_excel():
    """从 Excel 导入布局"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    f = request.files.get('file')
    if not f:
        return jsonify({'error': '请上传文件'}), 400
    try:
        import openpyxl
        wb = openpyxl.load_workbook(f)
        ws = wb.active
        fields = []
        for row in ws.iter_rows(min_row=1, values_only=False):
            for cell in row:
                if cell.value and str(cell.value).strip():
                    fields.append({
                        'key': f'field_{cell.row}_{cell.column}',
                        'label': str(cell.value).strip(),
                        'x': cell.column - 1,
                        'y': cell.row - 1,
                        'w': 2,
                        'h': 1,
                        'align': 'left',
                        'font_size': 9,
                        'bold': False
                    })
        # 自动识别可映射字段
        field_map = {'维修日期': 'date', '日期': 'date', '科室': 'department', '部门': 'department',
                     '设备': 'device_type', '打印机': 'device_type', '型号': 'device_model',
                     '编号': 'asset_no', '资产编号': 'asset_no', '内容': 'repair_content',
                     '维修内容': 'repair_content', '金额': 'total_amount', '合计': 'total_amount',
                     '品牌': 'device_brand', '维修人员': 'repair_person', '工单编号': 'work_order_no',
                     '备注': 'remark', '验收人': 'acceptor', '签字': 'acceptor',
                     '完成': 'complete_status', '完成情况': 'complete_status'}
        for f in fields:
            for cn, key in field_map.items():
                if cn in f['label']:
                    f['key'] = key
                    break
            if f['key'].startswith('field_'):
                f['key'] = f'custom_{f["key"]}'
        layout = {
            'cols': 12, 'rows': max(24, ws.max_row + 2),
            'page_width': 196, 'page_height': 277,
            'fields': fields,
            'title': '维修验收单',
            'title_font_size': 12,
            'border_style': 'solid',
            'border_width': 1
        }
        return jsonify({'ok': True, 'layout': layout})
    except Exception as e:
        return jsonify({'error': f'解析失败: {str(e)}'}), 400


@finance_bp.route('/acceptance_designer/import_word', methods=['POST'])
@login_required
def acceptance_designer_import_word():
    """从 Word 导入布局"""
    if not check_access():
        return jsonify({'error': '无权限'}), 403
    f = request.files.get('file')
    if not f:
        return jsonify({'error': '请上传文件'}), 400
    try:
        from docx import Document
        doc = Document(f)
        fields = []
        y = 0
        # 从段落提取
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                fields.append({
                    'key': f'para_{y}',
                    'label': text,
                    'x': 0,
                    'y': y,
                    'w': 12,
                    'h': 1,
                    'align': 'left',
                    'font_size': 9,
                    'bold': False
                })
                y += 1
        # 从表格提取
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        fields.append({
                            'key': f'table_{y}_{col_idx}',
                            'label': text,
                            'x': col_idx,
                            'y': y + row_idx,
                            'w': 2,
                            'h': 1,
                            'align': 'left',
                            'font_size': 9,
                            'bold': False
                        })
            y += len(table.rows)
        # 自动映射
        field_map = {'维修日期': 'date', '日期': 'date', '科室': 'department', '部门': 'department',
                     '设备': 'device_type', '打印机': 'device_type', '型号': 'device_model',
                     '编号': 'asset_no', '资产编号': 'asset_no', '内容': 'repair_content',
                     '维修内容': 'repair_content', '金额': 'total_amount', '合计': 'total_amount',
                     '品牌': 'device_brand', '维修人员': 'repair_person', '工单编号': 'work_order_no',
                     '备注': 'remark', '验收人': 'acceptor', '签字': 'acceptor',
                     '完成': 'complete_status', '完成情况': 'complete_status'}
        for f in fields:
            for cn, key in field_map.items():
                if cn in f['label']:
                    f['key'] = key
                    break
            if f['key'].startswith('table_') or f['key'].startswith('para_'):
                f['key'] = f'custom_{f["key"]}'
        layout = {
            'cols': 12, 'rows': max(24, y + 2),
            'page_width': 196, 'page_height': 277,
            'fields': fields,
            'title': '维修验收单',
            'title_font_size': 12,
            'border_style': 'solid',
            'border_width': 1
        }
        return jsonify({'ok': True, 'layout': layout})
    except ImportError:
        return jsonify({'error': '服务器未安装 python-docx，无法解析 Word 文件'}), 400
    except Exception as e:
        return jsonify({'error': f'解析失败: {str(e)}'}), 400
