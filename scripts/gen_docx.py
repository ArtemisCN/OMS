"""Convert 产品手册 to professional Word document (.docx)"""
import re, sys, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = '/home/ubuntu/hospital-workorder/产品手册_医院IT工单管理系统.md'
DOCX_PATH = '/home/ubuntu/hospital-workorder/static/demo/医院IT工单管理系统_操作手册.docx'

# ── helper: set cell shading ──
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

# ── read markdown ──
with open(MD_PATH, 'r', encoding='utf-8') as f:
    md = f.read()

doc = Document()

# ── styles ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35

# section margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

# ── helper: add styled paragraph ──
def add_para(text, bold=False, italic=False, size=10.5, color=None, align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_code_block(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    p._element.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '1E3A5F')
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_cell_shading(cell, 'F5F7FA')
            set_cell_text(cell, str(val), size=9)
    # column widths
    col_width = Inches(5.5 / len(headers))
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width
    doc.add_paragraph()  # spacing
    return table

# ── COVER PAGE ──
doc.add_paragraph()  # spacer
add_para('医院IT工单管理系统', bold=True, size=32,
         color=RGBColor(0x1e, 0x3a, 0x5f),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=120)
add_para('操 作 手 册', bold=True, size=20,
         color=RGBColor(0x47, 0x72, 0xa8),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=40)
add_para('版本 2.0 ｜ 2026年6月', size=12,
         color=RGBColor(0x66, 0x66, 0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('适用场景：医院信息科 / IT运维部门 / 后勤保障中心', size=11,
         color=RGBColor(0x88, 0x88, 0x88),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
add_para('本文档为内部使用，未经授权不得外传', size=9,
         color=RGBColor(0xaa, 0xaa, 0xaa),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=30)
doc.add_page_break()

# ── TABLE OF CONTENTS (manual) ──
add_para('目录', bold=True, size=18, color=RGBColor(0x1e, 0x3a, 0x5f),
         space_before=12, space_after=12)
toc_items = [
    ('1', '产品概述'),
    ('2', '系统架构'),
    ('3', '功能模块详解'),
    ('3.1', '仪表盘（Dashboard）'),
    ('3.2', '工单管理（Orders）'),
    ('3.3', 'Mobile Web 端（接单端）'),
    ('3.4', '微信小程序端'),
    ('3.5', 'Android 接单端'),
    ('3.6', '数据管理'),
    ('3.7', '资产台账'),
    ('3.8', '备件库存'),
    ('3.9', '耗材管理'),
    ('3.10', '巡检管理'),
    ('3.11', '电子表单'),
    ('3.12', '维修管理'),
    ('3.13', '值班排班'),
    ('3.14', '月度报表'),
    ('3.15', '审计日志'),
    ('3.16', '权限管理'),
    ('4', '三端协同工作机制'),
    ('5', '部署指南'),
    ('6', '用户角色与权限体系'),
    ('7', '运维管理'),
    ('8', '常见问题与排障'),
    ('9', '附录'),
]
for num, title in toc_items:
    indent = '    ' if '.' in num else ''
    p = doc.add_paragraph()
    r = p.add_run(f'{indent}{num}  {title}')
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(10.5 if '.' not in num else 10)
    if '.' not in num:
        r.bold = True
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── parse markdown into sections ──
lines = md.split('\n')
i = 0
in_table = False
table_headers = []
table_rows = []
table_mode = False
in_code_block = False
code_lines = []

def flush_table():
    global table_headers, table_rows
    if table_headers and table_rows:
        add_table(table_headers, table_rows)
    table_headers = []
    table_rows = []

def flush_code():
    global in_code_block, code_lines
    if code_lines:
        add_code_block('\n'.join(code_lines))
        code_lines = []

while i < len(lines):
    line = lines[i]
    
    # ── skip frontmatter before first heading ──
    if i < 20 and line.startswith('>') and '#' not in md[:500]:
        i += 1
        continue

    # ── code block ──
    if line.startswith('```'):
        if in_code_block:
            flush_code()
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # ── tables ──
    if '|' in line and line.strip().startswith('|'):
        cols = [c.strip() for c in line.split('|') if c.strip()]
        if not cols:
            i += 1
            continue
        # skip separator rows like |---|---|
        if all('-' in c for c in cols):
            table_mode = True
            i += 1
            continue
        if not table_mode:
            table_headers = cols
            table_mode = True
        else:
            table_rows.append(cols)
        i += 1
        continue
    else:
        if table_mode and table_headers:
            flush_table()
        table_mode = False

    # ── headings ──
    if line.startswith('## '):
        add_para(line[3:].strip(), bold=True, size=14,
                 color=RGBColor(0x1e, 0x3a, 0x5f),
                 space_before=18, space_after=8)
    elif line.startswith('### '):
        add_para(line[4:].strip(), bold=True, size=12,
                 color=RGBColor(0x2d, 0x5a, 0x87),
                 space_before=14, space_after=6)
    elif line.startswith('# '):
        # skip H1 (used for main title in MD)
        pass

    # ── hr ──
    elif line.strip() == '---':
        p = doc.add_paragraph()
        run = p.add_run('─' * 60)
        run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        run.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # ── list items ──
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        add_bullet(text)

    # ── numbered items ──
    elif re.match(r'^\d+[\.\、]', line.strip()):
        text = re.sub(r'^\d+[\.\、]\s*', '', line.strip())
        add_bullet(text)

    # ── blank line ──
    elif not line.strip():
        pass

    # ── regular paragraph ──
    else:
        text = line.strip()
        if text:
            # Handle bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

    i += 1

# flush remaining
if table_headers and table_rows:
    flush_table()

# ── ADD DIAGRAM AND ANIMATION REFERENCE PAGE ──
doc.add_page_break()
add_para('在线演示与可视化资源', bold=True, size=18,
         color=RGBColor(0x1e, 0x3a, 0x5f),
         space_before=12, space_after=12)
add_para('以下资源可通过浏览器直接访问，无需安装任何软件：', size=10.5, space_after=8)

add_table(
    ['资源', '链接', '说明'],
    [
        ['📐 系统架构图', 'https://demolin.cn/static/demo/architecture.html',
         '深色主题架构拓扑图，含三端前端、Nginx、Flask、数据库及外部通知集成'],
        ['🎬 工作流动画', 'https://demolin.cn/static/demo/animated-workflow.html',
         '工单生命周期、紧急程度流转、三端协同三大动画演示，纯CSS动画，可悬停暂停'],
        ['📘 本手册在线版', 'https://demolin.cn/static/demo/产品手册_医院IT工单管理系统.md',
         'Markdown 格式完整产品手册'],
        ['🎯 演示中心', 'https://demolin.cn/demo',
         '所有演示资源的统一入口页面（需登录后访问）'],
    ]
)

# ── SAVE ──
doc.save(DOCX_PATH)
print(f'OK: {DOCX_PATH}')
print(f'Size: {os.path.getsize(DOCX_PATH) / 1024:.0f} KB')
