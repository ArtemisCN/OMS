"""生成带截图的 Word 操作手册 - 图文并茂版"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCREENSHOTS = '/home/ubuntu/hospital-workorder/static/demo/screenshots'
OUTPUT = '/home/ubuntu/hospital-workorder/static/demo/医院IT工单管理系统_操作手册.docx'

doc = Document()

# ═══ styles ═══
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    if level == 1: hs.font.size = Pt(20)
    elif level == 2: hs.font.size = Pt(15)
    else: hs.font.size = Pt(12)

# ═══ helpers ═══
def add_para(text, bold=False, size=10.5, color=None, align=None, before=0, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def add_bullet(text, level=0, bold_prefix=None, size=10):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); r.font.size = Pt(size)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A5F"/>')
        c._tc.get_or_add_tcPr().append(s)
        c.text = ''; r = c.paragraphs[0].add_run(h); r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(9)
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            if ri % 2 == 1:
                s2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA"/>')
                c._tc.get_or_add_tcPr().append(s2)
            c.text = ''; r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9); r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

def add_step(num, title, desc, img, before=10):
    """Add a step with screenshot"""
    add_para(f'步骤{num}：{title}', bold=True, size=12, before=before, after=2)
    add_para(desc, size=10, after=4)
    img_path = os.path.join(SCREENSHOTS, img)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(img_path, width=Inches(5.5))
        add_para(f'▲ 图{num}：{title}', size=9, color=RGBColor(0x66,0x66,0x66),
                 align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=6)

def add_page_break():
    doc.add_page_break()

def add_tip(text, icon='💡'):
    add_para(f'{icon} {text}', size=9.5, color=RGBColor(0x47,0x72,0xa8), before=4, after=4)

# ═══════════════════════════ COVER ═══════════════════════════
for _ in range(4): doc.add_paragraph()
add_para('医院IT工单管理系统', bold=True, size=32,
         color=RGBColor(0x1e,0x3a,0x5f), align=WD_ALIGN_PARAGRAPH.CENTER, before=60)
add_para('操 作 手 册', bold=True, size=22,
         color=RGBColor(0x47,0x72,0xa8), align=WD_ALIGN_PARAGRAPH.CENTER, after=60)
add_para('版本 2.0 ｜ 2026年6月', size=12, color=RGBColor(0x66,0x66,0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('适用场景：医院信息科 / IT运维部门 / 后勤保障中心', size=11,
         color=RGBColor(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, before=6)
add_para('系统地址：https://demolin.cn', size=10, color=RGBColor(0x88,0x88,0x88),
         align=WD_ALIGN_PARAGRAPH.CENTER, before=4)
for _ in range(6): doc.add_paragraph()
add_para('本文档为内部使用，未经授权不得外传', size=9,
         color=RGBColor(0xaa,0xaa,0xaa), align=WD_ALIGN_PARAGRAPH.CENTER)
add_page_break()

# ═══════════════════════════ TOC ═══════════════════════════
add_para('目录', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), after=10)
toc = [
    ('1', '系统登录与首页'),
    ('2', '仪表盘（数据总览）'),
    ('3', '工单管理'),
    ('  3.1', '查看工单列表'),
    ('  3.2', '发布工单'),
    ('  3.3', '查看与编辑工单详情'),
    ('  3.4', '批量生成工单'),
    ('4', 'Mobile Web 端（手机接单）'),
    ('5', '微信小程序端'),
    ('6', '数据管理'),
    ('  6.1', '人员管理'),
    ('  6.2', '方案模板'),
    ('  6.3', '科室与地址管理'),
    ('  6.4', '知识库'),
    ('  6.5', '值班排班'),
    ('7', '资产台账'),
    ('8', '备件与耗材管理'),
    ('9', '巡检管理'),
    ('10', '电子表单'),
    ('11', '维修管理'),
    ('12', '月度报表'),
    ('13', '审计日志'),
    ('14', '权限管理'),
    ('15', '常见问题与排障'),
]
for num, title in toc:
    indent = title[:3] if '.' in title else ''
    display = f'  {num} {title}' if '.' in num else f'{num}  {title}'
    p = doc.add_paragraph()
    r = p.add_run(display)
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(10.5 if '.' not in num else 9.5)
    if '.' not in num: r.bold = True
    p.paragraph_format.space_after = Pt(2)
add_page_break()

# ═══════════════════════════ 1. LOGIN ═══════════════════════════
add_para('1  系统登录与首页', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_para('1.1  系统访问', bold=True, size=12, before=8)
add_bullet('PC端浏览器访问：https://demolin.cn', bold_prefix='地址：')
add_bullet('Mobile Web端：https://demolin.cn/mobile/', bold_prefix='手机端：')
add_bullet('微信小程序：搜索或扫码打开工单管理小程序', bold_prefix='小程序：')
add_bullet('推荐使用 Chrome / Edge 浏览器，分辨率建议 1366×768 以上', bold_prefix='浏览器：')

add_para('1.2  登录系统', bold=True, size=12, before=8)
add_step(1, '打开登录页', '在浏览器中打开系统地址，进入登录页面。输入用户名和密码后点击"登录"按钮。', '01-login.png')
add_tip('默认管理员账号：admin / admin123。首次登录后建议立即修改密码。')

add_para('1.3  首页概览', bold=True, size=12, before=8)
add_step(2, '仪表盘首页', '登录后自动进入仪表盘页面。页面按模块区域展示工单统计、故障分布、人员排行等数据。左侧为导航菜单，点击可进入各个功能模块。', '02-dashboard.png')
add_tip('仪表盘数据实时刷新，每次进入页面都会重新统计。所有卡片标题统一前缀"本月"。')
add_page_break()

# ═══════════════════════════ 2. DASHBOARD ═══════════════════════════
add_para('2  仪表盘（数据总览）', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_para('仪表盘是系统的数据中枢，展示本月运维全貌。以下为各区域说明：', size=10.5)

add_table(
    ['区域', '展示内容', '用途'],
    [
        ['本月统计卡', '工单总数、已完成、进行中、待接单（含环比）', '快速了解本月运维整体情况'],
        ['故障类型分布', '环形图展示各类故障占比', '分析高频故障类型，针对性优化'],
        ['人员排行', '横向柱状图展示工程师接单量 Top', '考核工作量，发现效率瓶颈'],
        ['近7天趋势', '折线图展示每日新建 vs 完成工单', '观察工单量走势，评估团队负载'],
        ['响应时长趋势', '折线图展示平均响应时长（分钟）', '监控服务响应效率'],
        ['今日动态', '滚动列表展示最新工单状态变化', '实时跟踪工单处理进展'],
        ['楼区热度', '横向柱状图各楼栋工单量排行', '定位高频报修区域'],
        ['科室排行', '横向柱状图各科室报修量排行', '识别重点服务科室'],
    ]
)

add_tip('图表区域支持轮播显示，每 7 秒自动切换，也可手动点击圆点指示器切换。响应时长 Y 轴自动缩放，上限 60 分钟。')
add_page_break()

# ═══════════════════════════ 3. ORDERS ═══════════════════════════
add_para('3  工单管理', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)
add_para('工单管理是系统的核心功能模块，包含工单列表、发布工单、批量生成等功能。', size=10.5)

# 3.1 Order List
add_para('3.1  查看工单列表', bold=True, size=14, before=10)
add_step(3, '工单列表页', '进入工单列表后，默认显示所有待接单工单。页面顶部有三个标签页：待接单（⏳）、处理中（🔧）、已完成（✅）。可以切换标签查看不同状态的工单。', '03-order-list.png')

add_para('筛选与搜索：', bold=True, size=10.5, before=6)
add_bullet('填写日期范围，点击"筛选"按钮', bold_prefix='按日期：')
add_bullet('选择楼区/楼层/科室三级联动下拉', bold_prefix='按区域：')
add_bullet('选择设备类型或故障类型', bold_prefix='按类型：')
add_bullet('输入标题、位置或描述关键词', bold_prefix='按关键词：')
add_bullet('点击下载按钮导出当前筛选结果', bold_prefix='导出Excel：')

add_para('工单列表字段说明：', bold=True, size=10.5, before=6)
add_table(
    ['字段', '说明'],
    [
        ['优先级圆点', '🔴紧急 / 🟡加急 / 🟢普通。已完成工单显示创建时的原始优先级，冻结不变'],
        ['工单标题', '点击标题进入工单详情页'],
        ['设备/故障类型', '自动识别的设备类型和故障分类'],
        ['位置', '楼区-楼层-科室-详细位置'],
        ['状态', '待接单(蓝) / 处理中(橙) / 已完成(绿)'],
        ['处理人', '已接单工单显示当前处理工程师'],
        ['时间', '显示发布于多久前（自动换算分钟/小时/天）'],
    ]
)

add_tip('排序规则：紧急>加急>普通，同级按时间倒序。')
add_page_break()

# 3.2 Publish Order
add_para('3.2  发布工单', bold=True, size=14, before=10)
add_step(4, '发布工单页', '在侧边栏点击"发布工单"进入发布页面。填写工单信息后点击"发布"按钮即可发布。系统会自动识别设备类型和故障类型。', '04-publish-form.png')

add_para('填写说明：', bold=True, size=10.5, before=6)
add_bullet('输入简洁明了的工单标题（如"3楼护士站打印机卡纸"）', bold_prefix='工单标题：')
add_bullet('系统根据标题关键词自动匹配（打印机→打印机，断网→网络设备）', bold_prefix='设备类型：')
add_bullet('系统根据标题自动分类（卡纸→打印机故障，蓝屏→硬件）', bold_prefix='故障类型：')
add_bullet('选择楼区→自动加载楼层→选择科室→自动填入楼栋位置', bold_prefix='位置信息：')
add_bullet('输入标题后系统会自动推荐匹配的维修方案', bold_prefix='方案模板：')
add_bullet('普通🟢 / 加急🟡 / 紧急🔴', bold_prefix='紧急程度：')
add_bullet('发布后自动推送企业微信群和微信订阅消息', bold_prefix='自动通知：')

add_tip('发布工单后，待接单标签页中会立即出现新的工单卡片。手机端和小程序端会收到通知提醒。')
add_page_break()

# 3.3 Order Detail & Edit
add_para('3.3  查看与编辑工单详情', bold=True, size=14, before=10)
add_step(5, '工单详情页', '在工单列表中点击工单标题进入详情页。页面展示工单的完整信息，包括发布时间线、位置、设备信息、处理方案等。', '05-order-detail.png')

add_para('详情页功能：', bold=True, size=10.5, before=6)
add_bullet('查看工单标题、设备类型、故障类型、紧急程度', bold_prefix='信息查看：')
add_bullet('完整的地址信息（楼区→楼层→科室→位置）', bold_prefix='位置查看：')
add_bullet('发布时间、接单时间、完成时间（时间线展示）', bold_prefix='时间线：')
add_bullet('点击"编辑"按钮可修改工单信息（已结单不可改紧急程度）', bold_prefix='编辑工单：')
add_bullet('管理员可删除工单（操作记录写入审计日志）', bold_prefix='删除工单：')
add_bullet('可循环切换优先级排序：普通→加急→紧急→普通', bold_prefix='切换优先级：')
add_tip('已完成工单的紧急程度冻结在创建时的原始值，不允许再修改。', icon='🚫')

# 3.4 Batch
add_para('3.4  批量生成工单', bold=True, size=14, before=10)
add_para('支持一次性批量生成多条测试工单，方便演示或压力测试。', size=10)
add_bullet('在侧边栏点击"批量生成"', bold_prefix='进入：')
add_bullet('输入需要生成的数量（1-50条），点击"预览"', bold_prefix='预览：')
add_bullet('系统自动生成预览列表，确认无误后点击"确认生成"', bold_prefix='确认：')
add_bullet('生成后 5 分钟内可点击"撤回"撤销最近一次批量操作', bold_prefix='撤回：')

add_page_break()

# ═══════════════════════════ 4. MOBILE WEB ═══════════════════════════
add_para('4  Mobile Web 端（手机接单）', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)
add_para('Mobile Web 端是专为运维人员在手机上设计的接单工作台，与 PC 端使用同一账号登录、共享数据。', size=10.5)

add_step(6, '手机端工单列表', '在手机浏览器打开 https://demolin.cn/mobile/，使用 PC 端相同的账号密码登录。页面同样有三个标签页：待接单 / 处理中 / 今日已完成。顶部显示四项统计数字。', '06-mobile-list.png')

add_step(7, '手机端工单详情', '点击工单进入详情页，可查看完整工单信息。根据工单状态，底部会显示对应的操作按钮。', '07-mobile-detail.png')

add_para('手机端操作说明：', bold=True, size=10.5, before=6)
add_table(
    ['操作', '功能', '说明'],
    [
        ['接单', '点击"接单"按钮', '认领工单，状态变为"处理中"，person 设为当前用户'],
        ['填方案', '填写处理方案后提交', '工单完成，状态变为"已完成"'],
        ['一键结单', '点击"一键结单"按钮', '系统自动匹配最佳方案模板并完成工单'],
        ['巡检提交', '勾选巡检项目+签名', '提交巡检结果，状态变为"已完成"'],
        ['今日总结', '点击"今日总结"', '生成今日已完成工单总结文本，自动复制到剪贴板'],
    ]
)
add_tip('一键结单时，系统会根据工单标题自动匹配最合适的方案模板。匹配逻辑基于 config.py 中的 SOLUTION_TEMPLATES。', icon='⚡')
add_page_break()

# ═══════════════════════════ 5. WECHAT ═══════════════════════════
add_para('5  微信小程序端', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)
add_para('微信小程序提供了与 Mobile Web 端相同的功能，支持更便捷的微信通知和离线消息提醒。', size=10.5)

add_para('5.1  登录方式', bold=True, size=12, before=8)
add_bullet('打开小程序后进入登录页，输入与 PC 端相同的用户名和密码')
add_bullet('登录成功后自动尝试绑定微信号（微信自动登录，无需再次输密码）')
add_bullet('支持 wx.login 自动登录（需先在系统后台绑定微信 openid）')

add_para('5.2  工单处理', bold=True, size=12, before=8)
add_bullet('三标签切换：未接单 / 已接单 / 已完成', bold_prefix='工单列表：')
add_bullet('每 10 秒自动刷新待接单列表，新工单到达有 toast 提示+震动', bold_prefix='新单提醒：')
add_bullet('点击工单进入详情，支持接单、填方案、一键结单、巡检提交', bold_prefix='操作：')
add_bullet('自动匹配方案模板，快速完成工单', bold_prefix='一键结单：')
add_bullet('生成并复制今日工作总结文本', bold_prefix='今日总结：')
add_bullet('授权接收新工单订阅消息', bold_prefix='订阅通知：')

add_para('5.3  电子表单操作', bold=True, size=12, before=8)
add_bullet('支持 10 种字段类型：文本、数字、邮箱、电话、下拉选择、单选、日期、复选框、手写签名')
add_bullet('内联编辑，不跳转页面，800ms 自动保存防抖')
add_bullet('手写签名使用 Canvas 2D 触摸绘制')
add_bullet('表单填写完成后提交审批，审批通过后自动完结关联工单')

add_tip('小程序端表单字段的值在 WXML 中预计算（JS 层），避免模板中直接使用方法调用导致编译错误。', icon='📝')
add_page_break()

# ═══════════════════════════ 6. DATA MANAGEMENT ═══════════════════════════
add_para('6  数据管理', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)
add_para('数据管理模块包含 8 个子模块，用于维护系统的基础数据。仅管理员或有权限的用户可访问。', size=10.5)

add_step(8, '数据管理首页', '在侧边栏点击"数据管理"进入总览页。页面展示各模块的数据统计卡片，点击即可进入对应的管理页面。', '08-data-manage.png')

# 6.1 Persons
add_para('6.1  人员管理', bold=True, size=14, before=10)
add_step(9, '人员管理列表', '在数据管理首页点击"人员管理"进入。展示所有运维工程师名单，支持启用/停用、编辑信息、创建登录账号等操作。', '09-persons.png')

add_para('操作说明：', bold=True, size=10.5, before=6)
add_bullet('点击"新增人员"添加新的工程师', bold_prefix='新增：')
add_bullet('点击"编辑"按钮修改电话、组别、备注', bold_prefix='编辑：')
add_bullet('点击开关切换启用/停用状态', bold_prefix='启用/停用：')
add_bullet('点击"创建账号"为该人员创建系统登录账号', bold_prefix='创建账号：')
add_bullet('自动从已有工单提取处理人姓名', bold_prefix='从工单导入：')

# 6.2 Solutions
add_para('6.2  方案模板', bold=True, size=14, before=10)
add_step(10, '方案模板列表', '在数据管理首页点击"方案模板"进入。系统预置 200+ 条维修方案模板，覆盖电脑、打印机、网络、软件等常见故障。支持搜索、编辑、新增。', '10-solutions.png')

add_para('操作说明：', bold=True, size=10.5, before=6)
add_bullet('搜索框输入关键词快速定位匹配方案', bold_prefix='搜索：')
add_bullet('点击"编辑"修改方案标题或内容', bold_prefix='编辑：')
add_bullet('填写方案标题和内容后保存', bold_prefix='新增：')
add_bullet('一键恢复为系统出厂默认值（覆盖所有自定义修改）', bold_prefix='重置默认：')
add_bullet('从已有工单的处理方案中提取并保存为模板', bold_prefix='从工单导入：')

# 6.3 Address
add_para('6.3  科室与地址管理', bold=True, size=14, before=10)
add_para('地址数据：', bold=True, size=10.5, before=6)
add_bullet('全院地址树管理，工单发布时的位置数据来源')
add_bullet('支持覆盖新增、软删除、搜索筛选')
add_bullet('数据来源：services/address.py（700+ 条预置地址）')

add_para('科室字典：', bold=True, size=10.5, before=4)
add_bullet('科室名称、楼层、楼区、联系电话')
add_bullet('支持从工单和地址数据自动导入')

# 6.4 Knowledge
add_para('6.4  知识库', bold=True, size=14, before=10)
add_step(11, '知识库', '知识库用于存储和分享运维知识文章，支持分类筛选、新增、编辑、删除。', '17-knowledge.png')
add_bullet('按分类筛选查看不同类别的文章', bold_prefix='分类浏览：')
add_bullet('填写标题、分类和正文内容', bold_prefix='新增文章：')
add_bullet('支持 Markdown 格式的富文本内容', bold_prefix='内容格式：')

# 6.5 Duty
add_para('6.5  值班排班', bold=True, size=14, before=10)
add_step(12, '值班排班', '网格视图展示排班表：横向为日期，纵向为人员。支持按月切换、单元格编辑、批量填充等操作。', '18-duty-schedules.png')
add_bullet('点击单元格快速设置/清除班次', bold_prefix='单格编辑：')
add_bullet('填充整行、填充工作日、清空整行、复制上个月', bold_prefix='批量操作：')
add_bullet('通过 Excel 文件批量导入排班数据', bold_prefix='导入：')

add_page_break()

# ═══════════════════════════ 7. ASSETS ═══════════════════════════
add_para('7  资产台账', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_step(13, '资产保修日历', '进入资产台账后默认显示保修日历视图，按剩余保修天数排序。三种筛选：即将过期(30天内)、已过保、有效期内。', '12-asset-calendar.png')

add_step(14, '资产列表', '切换到列表视图，可搜索、筛选、分页查看全部资产信息。支持批量编辑、调拨、回收、移位操作。', '13-asset-list.png')

add_para('核心操作：', bold=True, size=10.5, before=6)
add_table(
    ['操作', '说明'],
    [
        ['新增资产', '填写 50+ 个字段：资产编号、品牌、型号、SN、CPU、内存、硬盘、系统、IP、MAC、位置、财务信息等'],
        ['编辑资产', '修改任意字段，变更记录写入操作日志'],
        ['批量编辑', '批量修改指定字段值（如统一更新科室）'],
        ['批量调拨', '批量变更科室归属'],
        ['批量回收', '批量标记为报废/回收状态'],
        ['批量移位', '批量变更存放位置（楼栋/楼层/位置）'],
        ['Excel导入', '下载导入模板后批量导入资产数据'],
        ['Excel导出', '导出当前筛选条件下的资产台账'],
        ['操作日志', '查看所有资产的变更历史记录'],
    ]
)

add_tip('资产保修状态自动计算：剩余天数 ≤ 0 → "已过保"；剩余天数 ≤ 30 → "即将过期"；其他 → "有效期内"。', icon='📅')
add_page_break()

# ═══════════════════════════ 8. STOCK ═══════════════════════════
add_para('8  备件与耗材管理', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_para('8.1  备件库存', bold=True, size=14, before=10)
add_step(15, '备件库存总览', '展示所有备件的库存情况。库存低于最低预警值时显示红色标记。支持分类筛选、新增、编辑、出入库操作。', '14-stock.png')

add_para('备件操作说明：', bold=True, size=10.5, before=6)
add_bullet('填写备件名称、型号、品牌、分类、库存数量等信息', bold_prefix='新增备件：')
add_bullet('修改备件信息（管理员操作）', bold_prefix='编辑备件：')
add_bullet('入库：选择备件、填写数量、备注。出库：填写数量、关联工单号、备注，自动更新库存余额', bold_prefix='出入库：')
add_bullet('支持下载导入模板后批量导入备件', bold_prefix='Excel导入：')

add_para('8.2  耗材管理', bold=True, size=14, before=10)
add_para('耗材管理与备件管理类似，支持：', size=10)
add_bullet('耗材列表搜索、库存预警', bold_prefix='列表：')
add_bullet('耗材名称、型号、品牌、分类、库存数量', bold_prefix='新增/编辑：')
add_bullet('入库/出库操作，自动更新库存', bold_prefix='出入库：')
add_bullet('下载表头模板后批量导入', bold_prefix='Excel导入：')

add_page_break()

# ═══════════════════════════ 9. INSPECTION ═══════════════════════════
add_para('9  巡检管理', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_step(16, '巡检模板列表', '巡检管理分为模板管理和计划管理两部分。模板预设巡检内容项（如"检查电源""检查温度""检查网络设备"等），计划将模板与位置、时间关联。', '15-inspection-templates.png')

add_step(17, '巡检计划列表', '发布巡检计划后，系统会在设定的时间自动生成巡检工单。工程师在手机端勾选巡检项目并手写签名后提交。', '16-inspection-plans.png')

add_para('操作流程：', bold=True, size=10.5, before=6)
add_table(
    ['步骤', '操作', '说明'],
    [
        ['1', '创建巡检模板', '预设巡检检查项列表（如：电源、温度、网络设备、消防等）'],
        ['2', '发布巡检计划', '选择模板、设定位置和时间。可设置计划自动检查，前端定时调用检查接口'],
        ['3', '到期生成工单', '到达设定时间后，系统自动创建巡检工单，推送至工程师手机端'],
        ['4', '现场执行巡检', '工程师在手机端逐项勾选 ✅❌⬜，手写签名确认'],
        ['5', '提交巡检结果', '提交后工单自动完成，可在 PC 端查看巡检结果详情'],
        ['6', '导出确认单', '巡检确认单可导出或打印，作为巡检记录归档'],
    ]
)

add_tip('巡检计划支持设定具体到期时间（如每个周一早上 8:00），到期自动触发。', icon='⏰')
add_page_break()

# ═══════════════════════════ 10. FORMS ═══════════════════════════
add_para('10  电子表单', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)
add_para('电子表单模块支持自定义表单模板，用于设备发放、维修验收等场景的电子化审批流程。', size=10.5)

add_step(18, '表单模板列表', '在侧边栏点击"电子表单"进入。先创建表单模板（定义字段类型和布局），再使用模板创建表单实例。', '22-form-templates.png')

add_para('支持的字段类型：', bold=True, size=10.5, before=6)
add_table(
    ['字段类型', '说明'],
    [
        ['text', '单行文本输入框'],
        ['number', '数字输入框'],
        ['email', '邮箱输入'],
        ['tel', '电话号码输入'],
        ['url', '网址输入'],
        ['textarea', '多行文本'],
        ['select', '下拉选择（支持动态数据源：科室/人员/位置）'],
        ['radio', '单选按钮'],
        ['date', '日期选择'],
        ['checkbox', '多选框'],
        ['signature', '手写签名（Canvas 2D 触摸）'],
        ['richtext', '富文本显示区域（Header/Divider/Label）'],
    ]
)

add_para('审批流程：', bold=True, size=10.5, before=6)
add_bullet('管理员定义表单模板（字段列表+布局）', bold_prefix='1. 模板创建：')
add_bullet('使用模板创建表单实例，填充数据', bold_prefix='2. 表单创建：')
add_bullet('草稿→发布（自动创建关联工单）→提交审批→审批通过（自动完结工单）', bold_prefix='3. 生命周期：')

add_tip('电子表单支持 A4 画布定位布局，可直接打印生成纸质表单。字段位置支持精确的 x/y/w/h 坐标定位。', icon='🖨️')
add_page_break()

# ═══════════════════════════ 11. REPAIR ═══════════════════════════
add_para('11  维修管理', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)
add_para('维修管理模块用于管理设备维修工单，支持完整的创建→提交→审批→完成的流程。', size=10.5)

add_bullet('按关键词、日期范围、状态筛选', bold_prefix='维修单列表：')
add_bullet('选择维修模板创建维修单，填写故障描述和位置', bold_prefix='创建维修单：')
add_bullet('查看维修单详情，支持打印、保存字段值', bold_prefix='查看详情：')
add_bullet('提交审批后，管理员可审批通过或驳回', bold_prefix='审批流程：')
add_bullet('支持手写签名（Canvas 2D）', bold_prefix='签名：')

add_page_break()

# ═══════════════════════════ 12. REPORT ═══════════════════════════
add_para('12  月度报表', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_step(19, '月度报表', '报表页面展示当月工单数据的 7 大维度统计图表：工单总量趋势、故障类型分布、人员工作量、科室分布、响应时效、完成率走势、紧急程度分布。', '19-report.png')

add_table(
    ['维度', '展示方式', '说明'],
    [
        ['工单总量趋势', '全年月度柱状图', '月度对比分析运维工作量变化'],
        ['故障类型分布', '环形图', '分析各类故障占比，指导资源配置'],
        ['人员工作量', '横向柱状图', '考核工程师月度工作量'],
        ['科室分布', '横向柱状图', '各科室报修量排行'],
        ['响应时效', '折线图', '每月平均响应时长变化趋势'],
        ['完成率走势', '折线图', '月度完成率变化趋势'],
        ['紧急程度分布', '柱状图', '各紧急级别工单的数量分布'],
    ]
)

add_tip('点击"下载报表"按钮，系统自动生成包含 7 个工作表的 Excel 文件，每个维度对应一个独立工作表。', icon='📊')
add_page_break()

# ═══════════════════════════ 13. AUDIT ═══════════════════════════
add_para('13  审计日志', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_step(20, '审计日志', '审计日志记录所有用户的操作行为，支持分页查看、按操作人和操作类型筛选。适用于安全审计和问题追溯。', '20-audit-logs.png')

add_table(
    ['字段', '说明'],
    [
        ['操作时间', '操作发生的具体时间'],
        ['操作人', '执行操作的用户'],
        ['操作类型', 'create/update/delete/login/logout'],
        ['目标类型', '操作的对象类型（work_order/user/person 等）'],
        ['目标描述', '操作的简要描述（如"删除工单#123"）'],
        ['详情', '操作的详细内容'],
    ]
)

add_para('审计统计：', bold=True, size=10.5, before=6)
add_bullet('今日操作总量统计', bold_prefix='今日统计：')
add_bullet('近期操作量趋势', bold_prefix='趋势：')
add_bullet('操作频率最高的用户排行', bold_prefix='TOP操作人：')

add_page_break()

# ═══════════════════════════ 14. PERMISSIONS ═══════════════════════════
add_para('14  权限管理', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_step(21, '权限管理', '权限管理页面配置各用户角色对各业务模块的访问权限。左侧为管理员权限配置，右侧为普通用户权限配置。', '11-permissions.png')

add_para('角色体系：', bold=True, size=10.5, before=6)
add_table(
    ['角色', '权限', '典型用户'],
    [
        ['系统管理员', '所有模块可见，可管理所有数据', '信息科主任、系统管理员'],
        ['普通用户', '按模块权限矩阵控制可见性', '一线IT运维人员、工单发布员'],
    ]
)

add_bullet('勾选各模块的"可见"复选框，保存即可生效', bold_prefix='配置方法：')
add_bullet('开启/关闭用户的管理员权限（管理员全开所有模块）', bold_prefix='设为管理员：')
add_bullet('从人员名单中同步生成系统登录账号', bold_prefix='同步用户：')

add_tip('权限配置实时生效，修改后用户无需重新登录即可看到变化。侧边栏菜单根据权限动态渲染。', icon='🛡️')
add_page_break()

# ═══════════════════════════ 15. FAQ ═══════════════════════════
add_para('15  常见问题与排障', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

faqs = [
    ('Q1: 已完成的工单为什么全是绿色？',
     '已完成工单显示 original_priority（创建时的原始紧急程度）。如果全部显示绿色，说明工单创建时都是普通级别，自动升级脚本只影响未接单阶段的显示。'),
    ('Q2: 未接单的工单会自动变红吗？',
     '会。自动升级脚本每5分钟扫描未接单工单，将普通→加急→紧急逐级升级。这是为了让长时间无人接单的工单获得关注。'),
    ('Q3: 已完成的工单还能修改紧急程度吗？',
     '不能。已完成工单的紧急程度冻结在创建时的原始值，点击优先级圆点会返回403错误。'),
    ('Q4: 企业微信通知收不到？',
     '检查 .env 中 WECOM_WEBHOOK_URL 配置是否正确。测试方法：curl -X POST -H "Content-Type: application/json" -d \'{"msgtype":"text","text":{"content":"测试消息"}}\' <你的WEBHOOK_URL>'),
    ('Q5: 微信订阅消息推送失败 errcode 43101？',
     '43101 表示该用户未授权订阅。需要在调用 wx.requestSubscribeMessage 弹窗时让用户点击"允许"。一次性订阅每次推送都需用户授权。'),
    ('Q6: 小程序编译报错 "Bad attr value"？',
     'WXML 模板中不能直接调用 .substring() .startsWith() .indexOf() 等方法。需要在 JS 层预计算后再传入模板。'),
    ('Q7: 小程序登录后返回 401？',
     '检查 config.js 中 API_BASE_URL 是否为 HTTPS 地址；检查小程序后台 request合法域名是否已添加；检查 Token 是否正确存储和携带。'),
    ('Q8: 页面访问 404 或 502？',
     '404：确认 nginx 配置中 server_name 和 proxy_pass 正确。502：确认 Gunicorn 进程正在运行。查看 nginx 错误日志：tail -f /var/log/nginx/error.log'),
    ('Q9: 导出报表/Excel 无数据？',
     '检查当月是否有工单数据。在 PC 端工单列表中确认当前筛选条件下有数据。报表数据依赖当月统计数据。'),
    ('Q10: 数据库备份怎么做？',
     '建议保留 3 份备份轮换：cp workorders.db workorders-$(date +%Y%m%d).db。备份前执行 PRAGMA wal_checkpoint 确保 WAL 日志已写入主库。'),
]

for q, a in faqs:
    add_para(q, bold=True, size=10.5, color=RGBColor(0x1e,0x3a,0x5f), before=8, after=2)
    add_para(a, size=10, after=4)

add_page_break()

# ═══ APPENDIX ═══
add_para('附录', bold=True, size=18, color=RGBColor(0x1e,0x3a,0x5f), before=6)

add_para('A. 演示资源', bold=True, size=14, before=8)
add_table(
    ['资源', '链接', '说明'],
    [
        ['📐 系统架构图', '/static/demo/architecture.html', '深色主题架构拓扑图'],
        ['🎬 工作流动画', '/static/demo/animated-workflow.html', '工单生命周期动画演示'],
        ['📘 本手册在线版', '/static/demo/产品手册_医院IT工单管理系统.md', 'Markdown 格式完整手册'],
        ['🎯 演示中心', '/demo', '所有演示资源统一入口'],
    ]
)

add_para('B. 常用默认账号', bold=True, size=14, before=10)
add_table(
    ['账号', '密码', '角色', '说明'],
    [
        ['admin', 'admin123', '管理员', '首次登录后建议修改密码'],
    ]
)

add_para('C. 系统架构速览', bold=True, size=14, before=10)
add_table(
    ['层级', '技术', '说明'],
    [
        ['前端（PC）', 'Flask + Jinja2 + Bootstrap 5', '管理后台'],
        ['前端（Mobile）', 'Flask + 响应式 HTML', '手机接单端'],
        ['前端（小程序）', '微信原生 WXML/WXSS/JS', '移动查单处理'],
        ['后端', 'Flask 3.0 + Python 3.11', '应用逻辑'],
        ['数据库', 'SQLite 3', '单文件部署'],
        ['服务器', 'Gunicorn + Nginx', '生产部署'],
        ['通知', '企业微信 + 微信订阅消息', '推送通知'],
    ]
)

add_para('D. 部署命令速查', bold=True, size=14, before=10)
add_para('启动服务：', bold=True, size=10.5, before=4)
code = 'cd /var/www/hospital-workorder && source venv/bin/activate && gunicorn -w 4 -b 127.0.0.1:5000 wsgi:app'
add_para(code, size=8.5, color=RGBColor(0x33,0x33,0x33))
add_para('重启服务：', bold=True, size=10.5, before=4)
add_para('sudo systemctl restart hospital-workorder', size=8.5, color=RGBColor(0x33,0x33,0x33))
add_para('查看日志：', bold=True, size=10.5, before=4)
add_para('journalctl -u hospital-workorder -f', size=8.5, color=RGBColor(0x33,0x33,0x33))

# ═══ SAVE ═══
doc.save(OUTPUT)
print(f'✓ 已生成: {OUTPUT}')
print(f'  大小: {os.path.getsize(OUTPUT)/1024:.0f} KB')
